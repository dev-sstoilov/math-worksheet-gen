import sys
import random
import json
import os
import tempfile
import math
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGroupBox,
    QLabel, QComboBox, QPushButton, QSpinBox, QTabWidget, QFileDialog,
    QMessageBox, QButtonGroup, QRadioButton, QTabBar, QCheckBox
)
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QFont, QColor, QPainter, QIcon   # Added QIcon
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml

# Define settings file path in temp directory
SETTINGS_FILE = os.path.join(tempfile.gettempdir(), "math_worksheet_settings.json")


class CheckBoxTabBar(QTabBar):
    """Custom QTabBar that supports small checkboxes on tabs (right side)
    and draws each tab with a static per-tab color provided via `tab_colors`.
    The checkboxes do NOT disable tab clicks — they only reflect/select which operations are enabled.

    This class also ensures any left-side tab buttons are removed to avoid stray UI artifacts.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.checkboxes = {}
        # mapping index -> color string (e.g. '#e8f8e8')
        self.tab_colors = {}
        # ensure default stylesheet doesn't draw tab background so our custom painting shows
        # also remove left padding that some styles reserve for a left-side widget
        self.setStyleSheet("QTabBar::tab { background: transparent; margin-left: 0px; padding-left: 0px; }")
        # Avoid any default left-side buttons or indicators
        self.setUsesScrollButtons(False)
        self.setExpanding(False)

        # Keep left-side slot cleared whenever tabs are added or when widget shows.
        # Some platform/styles re-create the left-side placeholder after construction; handle that.
        try:
            self.tabInserted.connect(self._on_tab_inserted)
        except Exception:
            pass

    def showEvent(self, event):
        """Ensure left-side slot is cleared when the bar is shown (final opportunity)."""
        super().showEvent(event)
        for i in range(self.count()):
            try:
                self.setTabButton(i, QTabBar.LeftSide, None)
            except Exception:
                pass

    def _on_tab_inserted(self, index):
        """Clear left-side slot for newly inserted tabs."""
        try:
            self.setTabButton(index, QTabBar.LeftSide, None)
        except Exception:
            pass

    def add_checkbox_to_tab(self, index, checked=True):
        """Attach a small checkbox widget to an existing tab index (right side)."""
        cb = QCheckBox()
        cb.setChecked(checked)
        cb.setToolTip(self.tabText(index))
        cb.setStyleSheet("margin-right:6px; margin-left:4px;")
        cb.setFixedSize(cb.sizeHint().width() + 8, cb.sizeHint().height())
        # Ensure checkbox is placed on the right side only
        self.setTabButton(index, QTabBar.RightSide, cb)
        # Also explicitly clear LeftSide (best-effort)
        try:
            self.setTabButton(index, QTabBar.LeftSide, None)
        except Exception:
            pass
        self.checkboxes[index] = cb
        return cb

    def checkbox_for_tab(self, index):
        return self.checkboxes.get(index)

    def paintEvent(self, event):
        """Custom paint: fill each tab rect with its configured color (if any), then let the
        base class draw tab contents (text, buttons). We draw background first so text and
        checkbox widgets appear above it.
        """
        painter = QPainter(self)
        for i in range(self.count()):
            rect = self.tabRect(i)
            color = self.tab_colors.get(i)
            if color:
                painter.fillRect(rect, QColor(color))
        painter.end()
        # Now call base implementation to draw tab labels and buttons on top
        super().paintEvent(event)

class MathWorksheetGenerator(QMainWindow):
    def __init__(self):
        super().__init__()
        # Set program icon if logo exists
        if os.path.exists("./logo.ico"):
            self.setWindowIcon(QIcon("./logo.ico"))
        elif os.path.exists("C:\Program Files\Worksheet-Gen\logo.ico"):
            self.setWindowIcon(QIcon("C:\Program Files\Worksheet-Gen\logo.ico"))
        else:
            self.setWindowIcon(QIcon())  # Default icon
            
        self.setWindowTitle("Генератор на Математически Упражнения")
        self.setGeometry(100, 100, 900, 760)

        # Central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)

        # Header
        header = QLabel("Генератор на Математически Упражнения")
        header_font = QFont()
        header_font.setPointSize(30)
        header_font.setBold(True)
        header_font.setItalic(True)
        header.setFont(header_font)
        header.setAlignment(Qt.AlignCenter)
        header.setStyleSheet("margin: 18px 0; color: #2c3e50;")
        main_layout.addWidget(header)

        # Settings panel
        self.settings_group = QGroupBox("Настройки")
        settings_layout = QVBoxLayout(self.settings_group)
        main_layout.addWidget(self.settings_group)

        # Number format toggle
        format_group = QGroupBox("Формат на числата")
        format_layout = QHBoxLayout(format_group)
        self.word_format = QRadioButton("Думи")
        self.number_format = QRadioButton("Числа")
        self.number_format.setChecked(True)
        format_layout.addWidget(self.word_format)
        format_layout.addWidget(self.number_format)
        self.format_group = QButtonGroup()
        self.format_group.addButton(self.word_format)
        self.format_group.addButton(self.number_format)
        settings_layout.addWidget(format_group)

        # Number of equations
        equations_group = QGroupBox("Брой уравнения")
        equations_layout = QHBoxLayout(equations_group)
        self.equation_count = QSpinBox()
        self.equation_count.setRange(1, 1000)
        self.equation_count.setValue(30)
        equations_layout.addWidget(self.equation_count)
        settings_layout.addWidget(equations_group)

        # Operations group
        self.operations_group = QGroupBox("Операции")
        operations_layout = QVBoxLayout(self.operations_group)

        # Operation tabs with checkboxes in the tab bar (right side)
        self.operation_tabs = QTabWidget()
        self.custom_tab_bar = CheckBoxTabBar()
        self.operation_tabs.setTabBar(self.custom_tab_bar)

        # We'll put the tabs and the small colored area below them inside a container
        self.tabs_container = QWidget()
        tabs_container_layout = QVBoxLayout(self.tabs_container)
        tabs_container_layout.setContentsMargins(0, 0, 0, 0)
        tabs_container_layout.addWidget(self.operation_tabs)

        # This widget represents the area under the tabs but above the generate button.
        # We'll color this container together with the tab appearance when a tab is selected.
        self.tabs_colored_area = QWidget()
        self.tabs_colored_area.setMinimumHeight(20)
        tabs_container_layout.addWidget(self.tabs_colored_area)

        operations_layout.addWidget(self.tabs_container)

        # --- Addition tab ---
        addition_tab = QWidget()
        addition_layout = QVBoxLayout(addition_tab)

        addition_settings = QGroupBox("Настройки за събиране")
        addition_settings_layout = QVBoxLayout(addition_settings)

        # Operand 1 range
        operand1_group = QGroupBox("Първо число")
        operand1_layout = QHBoxLayout(operand1_group)
        self.add_op1_min = QSpinBox()
        self.add_op1_min.setRange(1, 1000)
        self.add_op1_min.setValue(1)
        self.add_op1_max = QSpinBox()
        self.add_op1_max.setRange(1, 1000)
        self.add_op1_max.setValue(100)
        operand1_layout.addWidget(QLabel("От:"))
        operand1_layout.addWidget(self.add_op1_min)
        operand1_layout.addWidget(QLabel("До:"))
        operand1_layout.addWidget(self.add_op1_max)
        addition_settings_layout.addWidget(operand1_group)

        # Operand 2 range
        operand2_group = QGroupBox("Второ число")
        operand2_layout = QHBoxLayout(operand2_group)
        self.add_op2_min = QSpinBox()
        self.add_op2_min.setRange(1, 1000)
        self.add_op2_min.setValue(1)
        self.add_op2_max = QSpinBox()
        self.add_op2_max.setRange(1, 1000)
        self.add_op2_max.setValue(100)
        operand2_layout.addWidget(QLabel("От:"))
        operand2_layout.addWidget(self.add_op2_min)
        operand2_layout.addWidget(QLabel("До:"))
        operand2_layout.addWidget(self.add_op2_max)
        addition_settings_layout.addWidget(operand2_group)

        # Regrouping settings
        regrouping_group = QGroupBox("Настройки за преминаване")
        regrouping_layout = QVBoxLayout(regrouping_group)
        self.add_regrouping = QComboBox()
        self.add_regrouping.addItems(["Без преминаване", "Само с преминаване", "Смесени"])
        regrouping_layout.addWidget(self.add_regrouping)
        addition_settings_layout.addWidget(regrouping_group)

        addition_layout.addWidget(addition_settings)
        # Add tab and remember its index
        self.operation_tabs.addTab(addition_tab, "Събиране")
        self.add_tab_index = self.operation_tabs.indexOf(addition_tab)
        self.custom_tab_bar.add_checkbox_to_tab(self.add_tab_index, checked=True)

        # --- Subtraction tab ---
        subtraction_tab = QWidget()
        subtraction_layout = QVBoxLayout(subtraction_tab)

        subtraction_settings = QGroupBox("Настройки за изваждане")
        subtraction_settings_layout = QVBoxLayout(subtraction_settings)

        sub_op1_group = QGroupBox("Първо число")
        sub_op1_layout = QHBoxLayout(sub_op1_group)
        self.sub_op1_min = QSpinBox()
        self.sub_op1_min.setRange(1, 1000)
        self.sub_op1_min.setValue(1)
        self.sub_op1_max = QSpinBox()
        self.sub_op1_max.setRange(1, 1000)
        self.sub_op1_max.setValue(100)
        sub_op1_layout.addWidget(QLabel("От:"))
        sub_op1_layout.addWidget(self.sub_op1_min)
        sub_op1_layout.addWidget(QLabel("До:"))
        sub_op1_layout.addWidget(self.sub_op1_max)
        subtraction_settings_layout.addWidget(sub_op1_group)

        sub_op2_group = QGroupBox("Второ число")
        sub_op2_layout = QHBoxLayout(sub_op2_group)
        self.sub_op2_min = QSpinBox()
        self.sub_op2_min.setRange(1, 1000)
        self.sub_op2_min.setValue(1)
        self.sub_op2_max = QSpinBox()
        self.sub_op2_max.setRange(1, 1000)
        self.sub_op2_max.setValue(100)
        sub_op2_layout.addWidget(QLabel("От:"))
        sub_op2_layout.addWidget(self.sub_op2_min)
        sub_op2_layout.addWidget(QLabel("До:"))
        sub_op2_layout.addWidget(self.sub_op2_max)
        subtraction_settings_layout.addWidget(sub_op2_group)

        sub_regrouping_group = QGroupBox("Настройки за преминаване")
        sub_regrouping_layout = QVBoxLayout(sub_regrouping_group)
        self.sub_regrouping = QComboBox()
        self.sub_regrouping.addItems(["Без преминаване", "Само с преминаване", "Смесени"])
        sub_regrouping_layout.addWidget(self.sub_regrouping)
        subtraction_settings_layout.addWidget(sub_regrouping_group)

        subtraction_layout.addWidget(subtraction_settings)

        self.operation_tabs.addTab(subtraction_tab, "Изваждане")
        self.sub_tab_index = self.operation_tabs.indexOf(subtraction_tab)
        self.custom_tab_bar.add_checkbox_to_tab(self.sub_tab_index, checked=True)

        # --- Multiplication tab ---
        multiplication_tab = QWidget()
        multiplication_layout = QVBoxLayout(multiplication_tab)

        multiplication_settings = QGroupBox("Настройки за умножение")
        multiplication_settings_layout = QVBoxLayout(multiplication_settings)

        mul_op1_group = QGroupBox("Първо число")
        mul_op1_layout = QHBoxLayout(mul_op1_group)
        self.mul_op1_min = QSpinBox()
        self.mul_op1_min.setRange(1, 100)
        self.mul_op1_min.setValue(1)
        self.mul_op1_max = QSpinBox()
        self.mul_op1_max.setRange(1, 100)
        self.mul_op1_max.setValue(12)
        mul_op1_layout.addWidget(QLabel("От:"))
        mul_op1_layout.addWidget(self.mul_op1_min)
        mul_op1_layout.addWidget(QLabel("До:"))
        mul_op1_layout.addWidget(self.mul_op1_max)
        multiplication_settings_layout.addWidget(mul_op1_group)

        mul_op2_group = QGroupBox("Второ число")
        mul_op2_layout = QHBoxLayout(mul_op2_group)
        self.mul_op2_min = QSpinBox()
        self.mul_op2_min.setRange(1, 100)
        self.mul_op2_min.setValue(1)
        self.mul_op2_max = QSpinBox()
        self.mul_op2_max.setRange(1, 100)
        self.mul_op2_max.setValue(12)
        mul_op2_layout.addWidget(QLabel("От:"))
        mul_op2_layout.addWidget(self.mul_op2_min)
        mul_op2_layout.addWidget(QLabel("До:"))
        mul_op2_layout.addWidget(self.mul_op2_max)
        multiplication_settings_layout.addWidget(mul_op2_group)

        multiplication_layout.addWidget(multiplication_settings)

        self.operation_tabs.addTab(multiplication_tab, "Умножение")
        self.mul_tab_index = self.operation_tabs.indexOf(multiplication_tab)
        self.custom_tab_bar.add_checkbox_to_tab(self.mul_tab_index, checked=False)

        # --- Division tab ---
        division_tab = QWidget()
        division_layout = QVBoxLayout(division_tab)

        division_settings = QGroupBox("Настройки за деление")
        division_settings_layout = QVBoxLayout(division_settings)

        div_op1_group = QGroupBox("Първо число")
        div_op1_layout = QHBoxLayout(div_op1_group)
        self.div_op1_min = QSpinBox()
        self.div_op1_min.setRange(1, 1000)
        self.div_op1_min.setValue(1)
        self.div_op1_max = QSpinBox()
        self.div_op1_max.setRange(1, 1000)
        self.div_op1_max.setValue(144)
        div_op1_layout.addWidget(QLabel("От:"))
        div_op1_layout.addWidget(self.div_op1_min)
        div_op1_layout.addWidget(QLabel("До:"))
        div_op1_layout.addWidget(self.div_op1_max)
        division_settings_layout.addWidget(div_op1_group)

        div_op2_group = QGroupBox("Второ число")
        div_op2_layout = QHBoxLayout(div_op2_group)
        self.div_op2_min = QSpinBox()
        self.div_op2_min.setRange(1, 100)
        self.div_op2_min.setValue(1)
        self.div_op2_max = QSpinBox()
        self.div_op2_max.setRange(1, 100)
        self.div_op2_max.setValue(12)
        div_op2_layout.addWidget(QLabel("От:"))
        div_op2_layout.addWidget(self.div_op2_min)
        div_op2_layout.addWidget(QLabel("До:"))
        div_op2_layout.addWidget(self.div_op2_max)
        division_settings_layout.addWidget(div_op2_group)

        # New: allow non-integer division results (default ON)
        self.div_allow_decimal = QCheckBox("Разрешени десетични числа")
        self.div_allow_decimal.setChecked(True)
        division_settings_layout.addWidget(self.div_allow_decimal)

        division_layout.addWidget(division_settings)

        self.operation_tabs.addTab(division_tab, "Деление")
        self.div_tab_index = self.operation_tabs.indexOf(division_tab)
        self.custom_tab_bar.add_checkbox_to_tab(self.div_tab_index, checked=False)

        settings_layout.addWidget(self.operations_group)

        # Generate button
        generate_btn = QPushButton("Генерирай Упражнения + Отговори")
        generate_btn.clicked.connect(self.generate_documents)
        generate_btn.setStyleSheet(
            "background-color: #3498db; color: white; font-weight: bold; padding: 10px;"
        )
        settings_layout.addWidget(generate_btn)

        # Help button
        help_btn = QPushButton("Помощ")
        help_btn.clicked.connect(self.show_help)
        settings_layout.addWidget(help_btn)

        # Status bar
        self.statusBar().showMessage("Готов за генериране на упражнения")

        # Store generated problems and last directory
        self.problems = []
        self.last_dir = os.getcwd()

        # Tab colors map (distinct color per tab) — populate after tabs exist
        self._initial_tab_colors = {}

        # Connect tab change signal to update tabs and colored area
        self.operation_tabs.currentChanged.connect(self.on_tab_changed)

        # Apply style — make default text bigger
        self.setStyleSheet("""
            QWidget {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                font-size: 14pt;
            }
            QGroupBox {
                border: 1px solid #ddd;
                border-radius: 5px;
                margin-top: 1ex;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                padding: 0 5px;
                font-weight: bold;
            }
            QTabWidget::pane {
                border: 1px solid #ddd;
                padding: 10px;
            }
        """)

        # After adding tabs we can set up colors and ensure persistent settings are loaded
        self.setup_tab_colors()
        self.load_settings()

        # Set initial appearance according to currently selected tab
        self.on_tab_changed(self.operation_tabs.currentIndex())

    def setup_tab_colors(self):
        # populate mapping using tab indices (safe to call after tabs added)
        self._initial_tab_colors = {
            self.add_tab_index: '#e8f8e8',        # light green for addition
            self.sub_tab_index: '#fde6e6',        # light red/pink for subtraction
            self.mul_tab_index: '#fff5e6',        # light orange for multiplication
            self.div_tab_index: '#f3e8ff'         # light purple for division
        }
        # Give these colors to the custom tab bar so it paints each tab with its static color
        self.custom_tab_bar.tab_colors = self._initial_tab_colors.copy()
        # ensure base tabBar stylesheet doesn't override our colors
        self.operation_tabs.tabBar().setStyleSheet("QTabBar::tab { background: transparent; }")

        # IMPORTANT: remove any left-side tab widgets (these can appear as small rounded squares on some platforms)
        # This loop explicitly clears the LeftSide slot for every tab to remove that visual artifact.
        try:
            for i in range(self.custom_tab_bar.count()):
                self.custom_tab_bar.setTabButton(i, QTabBar.LeftSide, None)
        except Exception:
            # not critical; best-effort only
            pass

    # -------------------------
    # Settings persistence
    # -------------------------
    def save_settings(self):
        """Save current settings to a JSON file"""
        settings = {
            "number_format": "words" if self.word_format.isChecked() else "numbers",
            "equation_count": self.equation_count.value(),
            "last_dir": self.last_dir,
            "addition": {
                "enabled": bool(self.custom_tab_bar.checkbox_for_tab(self.add_tab_index).isChecked()),
                "op1_min": self.add_op1_min.value(),
                "op1_max": self.add_op1_max.value(),
                "op2_min": self.add_op2_min.value(),
                "op2_max": self.add_op2_max.value(),
                "regrouping": self.add_regrouping.currentIndex()
            },
            "subtraction": {
                "enabled": bool(self.custom_tab_bar.checkbox_for_tab(self.sub_tab_index).isChecked()),
                "op1_min": self.sub_op1_min.value(),
                "op1_max": self.sub_op1_max.value(),
                "op2_min": self.sub_op2_min.value(),
                "op2_max": self.sub_op2_max.value(),
                "regrouping": self.sub_regrouping.currentIndex()
            },
            "multiplication": {
                "enabled": bool(self.custom_tab_bar.checkbox_for_tab(self.mul_tab_index).isChecked()),
                "op1_min": self.mul_op1_min.value(),
                "op1_max": self.mul_op1_max.value(),
                "op2_min": self.mul_op2_min.value(),
                "op2_max": self.mul_op2_max.value()
            },
            "division": {
                "enabled": bool(self.custom_tab_bar.checkbox_for_tab(self.div_tab_index).isChecked()),
                "op1_min": self.div_op1_min.value(),
                "op1_max": self.div_op1_max.value(),
                "op2_min": self.div_op2_min.value(),
                "op2_max": self.div_op2_max.value(),
                "allow_decimal": bool(self.div_allow_decimal.isChecked())
            }
        }

        try:
            with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
                json.dump(settings, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self.statusBar().showMessage(f"Грешка при запис на настройки: {str(e)}", 5000)

    def load_settings(self):
        """Load settings from JSON file if exists"""
        if not os.path.exists(SETTINGS_FILE):
            return

        try:
            with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
                settings = json.load(f)
        except Exception as e:
            self.statusBar().showMessage(f"Грешка при зареждане на настройки: {str(e)}", 5000)
            return

        # Apply loaded settings
        if settings.get("number_format") == "words":
            self.word_format.setChecked(True)
        else:
            self.number_format.setChecked(True)

        self.equation_count.setValue(settings.get("equation_count", 30))

        # last_dir
        self.last_dir = settings.get("last_dir", os.getcwd())

        # Addition settings
        add_settings = settings.get("addition", {})
        add_enabled = add_settings.get("enabled", True)
        cb_add = self.custom_tab_bar.checkbox_for_tab(self.add_tab_index)
        if cb_add is not None:
            cb_add.setChecked(add_enabled)
        self.add_op1_min.setValue(add_settings.get("op1_min", 1))
        self.add_op1_max.setValue(add_settings.get("op1_max", 100))
        self.add_op2_min.setValue(add_settings.get("op2_min", 1))
        self.add_op2_max.setValue(add_settings.get("op2_max", 100))
        self.add_regrouping.setCurrentIndex(add_settings.get("regrouping", 0))

        # Subtraction settings
        sub_settings = settings.get("subtraction", {})
        sub_enabled = sub_settings.get("enabled", True)
        cb_sub = self.custom_tab_bar.checkbox_for_tab(self.sub_tab_index)
        if cb_sub is not None:
            cb_sub.setChecked(sub_enabled)
        self.sub_op1_min.setValue(sub_settings.get("op1_min", 1))
        self.sub_op1_max.setValue(sub_settings.get("op1_max", 100))
        self.sub_op2_min.setValue(sub_settings.get("op2_min", 1))
        self.sub_op2_max.setValue(sub_settings.get("op2_max", 100))
        self.sub_regrouping.setCurrentIndex(sub_settings.get("regrouping", 0))

        # Multiplication settings
        mul_settings = settings.get("multiplication", {})
        mul_enabled = mul_settings.get("enabled", False)
        cb_mul = self.custom_tab_bar.checkbox_for_tab(self.mul_tab_index)
        if cb_mul is not None:
            cb_mul.setChecked(mul_enabled)
        self.mul_op1_min.setValue(mul_settings.get("op1_min", 1))
        self.mul_op1_max.setValue(mul_settings.get("op1_max", 12))
        self.mul_op2_min.setValue(mul_settings.get("op2_min", 1))
        self.mul_op2_max.setValue(mul_settings.get("op2_max", 12))

        # Division settings
        div_settings = settings.get("division", {})
        div_enabled = div_settings.get("enabled", False)
        cb_div = self.custom_tab_bar.checkbox_for_tab(self.div_tab_index)
        if cb_div is not None:
            cb_div.setChecked(div_enabled)
        self.div_op1_min.setValue(div_settings.get("op1_min", 1))
        self.div_op1_max.setValue(div_settings.get("op1_max", 144))
        self.div_op2_min.setValue(div_settings.get("op2_min", 1))
        self.div_op2_max.setValue(div_settings.get("op2_max", 12))
        self.div_allow_decimal.setChecked(div_settings.get("allow_decimal", True))

    # -------------------------
    # Number to Bulgarian (unchanged for integers; floats return numeric string)
    # -------------------------
    def number_to_bulgarian(self, num):
        """Convert a number to Bulgarian words with correct grammar for integers.
        If the value is not an integer (e.g., decimals), return a numeric string representation.
        """
        if self.number_format.isChecked():
            if isinstance(num, float):
                if abs(num - round(num)) < 1e-9:
                    return str(int(round(num)))
                return f"{num:.1f}"
            return str(num)

        if not isinstance(num, int):
            if isinstance(num, float):
                return f"{num:.1f}"
            return str(num)

        units = ['', 'едно', 'две', 'три', 'четири', 'пет', 'шест', 'седем', 'осем', 'девет']
        teens = ['десет', 'единадесет', 'дванадесет', 'тринадесет', 'четиринадесет',
                 'петнадесет', 'шестнадесет', 'седемнадесет', 'осемнадесет', 'деветнадесет']
        tens = ['', '', 'двадесет', 'тридесет', 'четиридесет', 'петдесет',
                'шестдесет', 'седемдесет', 'осемдесет', 'деветдесет']
        hundreds = ['', 'сто', 'двеста', 'триста', 'четиристотин', 'петстотин',
                    'шестстотин', 'седемстотин', 'осемстотин', 'деветстотин']

        num = int(num)
        if num == 0:
            return 'нула'
        if num < 10:
            return units[num]
        if num < 20:
            return teens[num - 10]

        parts = []

        if num >= 100:
            hundreds_digit = num // 100
            parts.append(hundreds[hundreds_digit])
            num %= 100

        if num >= 20:
            tens_digit = num // 10
            parts.append(tens[tens_digit])
            num %= 10
            if num > 0:
                parts.append("и")
                parts.append(units[num])
        elif num >= 10:
            parts.append(teens[num - 10])
        elif num > 0:
            parts.append(units[num])

        return ' '.join(parts)

    def needs_regrouping(self, a, b, operation):
        """Determine if regrouping is needed for addition or subtraction"""
        if operation == "addition":
            return (int(a) % 10) + (int(b) % 10) >= 10
        elif operation == "subtraction":
            return (int(a) % 10) < (int(b) % 10)
        return False

    # -------------------------
    # Generators for operations (unchanged except division when decimals allowed)
    # -------------------------
    def generate_addition(self):
        min1 = self.add_op1_min.value()
        max1 = self.add_op1_max.value()
        min2 = self.add_op2_min.value()
        max2 = self.add_op2_max.value()
        regrouping_mode = self.add_regrouping.currentText()

        while True:
            a = random.randint(min1, max1)
            b = random.randint(min2, max2)
            result = a + b

            needs_regroup = self.needs_regrouping(a, b, "addition")

            if regrouping_mode == "Без преминаване" and needs_regroup:
                continue
            elif regrouping_mode == "Само с преминаване" and not needs_regroup:
                continue

            return a, b, result

    def generate_subtraction(self):
        min1 = self.sub_op1_min.value()
        max1 = self.sub_op1_max.value()
        min2 = self.sub_op2_min.value()
        max2 = self.sub_op2_max.value()
        regrouping_mode = self.sub_regrouping.currentText()

        while True:
            a = random.randint(min1, max1)
            b = random.randint(min2, max2)

            if a < b:
                a, b = b, a

            result = a - b
            needs_regroup = self.needs_regrouping(a, b, "subtraction")

            if regrouping_mode == "Без преминаване" and needs_regroup:
                continue
            elif regrouping_mode == "Само с преминаване" and not needs_regroup:
                continue

            return a, b, result

    def generate_multiplication(self):
        min1 = self.mul_op1_min.value()
        max1 = self.mul_op1_max.value()
        min2 = self.mul_op2_min.value()
        max2 = self.mul_op2_max.value()

        a = random.randint(min1, max1)
        b = random.randint(min2, max2)
        result = a * b

        return a, b, result

    def generate_division(self):
        min1 = self.div_op1_min.value()
        max1 = self.div_op1_max.value()
        min2 = self.div_op2_min.value()
        max2 = self.div_op2_max.value()

        allow_decimal = bool(self.div_allow_decimal.isChecked())

        # If decimals allowed, produce operands and quotient with one decimal place
        if allow_decimal:
            b_raw = random.randint(min2 * 10, max2 * 10)
            b = b_raw / 10.0
            min_q = max(0.1, min1 / b)
            max_q = max(0.1, max1 / b)
            if min_q > max_q:
                min_q, max_q = max_q, min_q
            q_raw = random.randint(int(math.floor(min_q * 10)), max(1, int(math.floor(max_q * 10))))
            quotient = q_raw / 10.0
            a = round(quotient * b, 1)
            result = quotient
            if a < min1 or a > max1:
                allow_decimal = False

        if not allow_decimal:
            b = random.randint(min2, max2)
            min_quotient = max(1, min1 // b)
            max_quotient = max1 // b

            if min_quotient > max_quotient:
                min_quotient, max_quotient = max_quotient, min_quotient

            quotient = random.randint(min_quotient, max_quotient)
            a = quotient * b
            result = quotient

        return a, b, result

    def is_operation_enabled(self, tab_index):
        cb = self.custom_tab_bar.checkbox_for_tab(tab_index)
        return bool(cb.isChecked()) if cb is not None else False

    def generate_problems(self, count):
        """Generate the specified number of math problems"""
        operations = []
        if self.is_operation_enabled(self.add_tab_index):
            operations.append("addition")
        if self.is_operation_enabled(self.sub_tab_index):
            operations.append("subtraction")
        if self.is_operation_enabled(self.mul_tab_index):
            operations.append("multiplication")
        if self.is_operation_enabled(self.div_tab_index):
            operations.append("division")

        if not operations:
            QMessageBox.warning(self, "Предупреждение", "Моля изберете поне една операция!")
            return []

        # Validate that count is reasonable
        if count <= 0:
            QMessageBox.warning(self, "Предупреждение", "Броят уравнения трябва да е поне 1!")
            return []

        problems = []
        used_combinations = set()
    
        # Safety counter to prevent infinite loops
        max_attempts = count * 100
        attempts = 0

        while len(problems) < count and attempts < max_attempts:
            attempts += 1
            op = random.choice(operations)

            try:
                if op == "addition":
                    a, b, result = self.generate_addition()
                    op_symbol = "+" if self.number_format.isChecked() else "плюс"
                elif op == "subtraction":
                    a, b, result = self.generate_subtraction()
                    op_symbol = "-" if self.number_format.isChecked() else "минус"
                elif op == "multiplication":
                    a, b, result = self.generate_multiplication()
                    op_symbol = "×" if self.number_format.isChecked() else "умножено по"
                else:  # division
                    a, b, result = self.generate_division()
                    op_symbol = "÷" if self.number_format.isChecked() else "делено на"
            except Exception as e:
                # Log the error but continue trying
                print(f"Error generating problem: {e}")
                continue

            key = f"{op}-{a}-{b}"
            if key in used_combinations:
                continue
            used_combinations.add(key)

            display_a = self.number_to_bulgarian(a)
            display_b = self.number_to_bulgarian(b)

            problem = {
                "question": f"{display_a} {op_symbol} {display_b} = ",
                "answer": f"{display_a} {op_symbol} {display_b} = {result}"
            }
            problems.append(problem)

        if len(problems) < count:
            self.statusBar().showMessage(f"Генерирани са само {len(problems)} от {count} уравнения (ограничени опции)", 5000)

        return problems[:count]

    def create_word_document(self, problems):
        """Create a Word document with the problems formatted appropriately

        Adjustments:
         - For numeric layout (numbers), create pages with 102 equations per page (6 columns x 17 rows = 102)
         - Align all equations (questions and answers) to the left
        """
        doc = Document()

        section = doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # If numeric table layout
        if self.number_format.isChecked():
            per_page = 102
            cols = 6

            # chunk problems into pages
            for page_start in range(0, len(problems), per_page):
                page_chunk = problems[page_start:page_start + per_page]

                # Calculate how many rows we need
                rows_needed = math.ceil(len(page_chunk) / cols)
                table = doc.add_table(rows=rows_needed, cols=cols)
            
                # Set column widths - do this after table is created but before populating
                for col in table.columns:
                    col.width = Inches(1.25)

                # remove borders for a cleaner look (keeps same behavior)
                try:
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    tblBorders = parse_xml(
                        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                        r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                        r'</w:tblBorders>'
                    )
                    tblPr.append(tblBorders)
                except Exception:
                    # If XML parsing fails, continue without custom borders
                    pass

                # Populate the table
                for i, problem in enumerate(page_chunk):
                    row_idx = i // cols
                    col_idx = i % cols
                
                    if row_idx < rows_needed and col_idx < cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                        # create left-aligned paragraph
                        p = cell.add_paragraph(problem)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.line_spacing = 1
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        for run in p.runs:
                            run.font.size = Pt(14)
                            run.font.name = 'Arial'

                # insert a page break after each page chunk except the last
                if page_start + per_page < len(problems):
                    doc.add_page_break()
        else:
            # For word-format or default behaviour, output one-per-line (left aligned)
            for problem in problems:
                p = doc.add_paragraph(problem)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.name = 'Arial'

        return doc

    def generate_documents(self):
        """Generate exercises and answers Word documents"""
        # Save settings before generating
        self.save_settings()

        count = self.equation_count.value()
        self.problems = self.generate_problems(count)

        if not self.problems:
            return

        exercises = [p["question"] for p in self.problems]
        exercises_doc = self.create_word_document(exercises)

        answers = [p["answer"] for p in self.problems]
        answers_doc = self.create_word_document(answers)

        try:
            # Use last_dir as the default directory for the save dialogs
            default_dir = self.last_dir if os.path.isdir(self.last_dir) else os.getcwd()
            default_ex_name = os.path.join(default_dir, "math_worksheets.docx")
            exercises_file, _ = QFileDialog.getSaveFileName(
                self, "Запази упражнения", default_ex_name, "Word Documents (*.docx)"
            )
            if exercises_file:
                exercises_doc.save(exercises_file)
                # update last_dir and persist
                self.last_dir = os.path.dirname(exercises_file)
                self.save_settings()

            default_ans_name = os.path.join(self.last_dir if os.path.isdir(self.last_dir) else default_dir, "math_answers.docx")
            answers_file, _ = QFileDialog.getSaveFileName(
                self, "Запази отговори", default_ans_name, "Word Documents (*.docx)"
            )
            if answers_file:
                answers_doc.save(answers_file)
                # update last_dir and persist
                self.last_dir = os.path.dirname(answers_file)
                self.save_settings()

            self.statusBar().showMessage(f"Успешно генерирани {count} уравнения", 5000)
            QMessageBox.information(self, "Готово", "Документите са успешно създадени!")
        except Exception as e:
            QMessageBox.critical(self, "Грешка", f"Възникна грешка при записването: {str(e)}")

    def show_help(self):
        """Show help dialog"""
        help_text = """
        <h3>Помощ за Генератора</h3>
        <p><strong>Формат на числата:</strong> Изберете дали уравненията да използват числа или думи</p>
        <p><strong>Настройки за преминаване:</strong></p>
        <ul>
            <li>Без: Няма преминаване</li>
            <li>Само: Само с преминаване</li>
            <li>Смесени: И двете смесени</li>
        </ul>
        <p>(Настройката е отделна за събиране и изваждане)</p>
        <p><strong>Персонализирани диапазони:</strong> За всяка операция можете да зададете отделни минимални и максимални стойности за всяко от числата в израза.</p>
        <p><strong>Брой уравнения:</strong> Въведете желания брой задачи (от 1 до 1000)</p>
        <p><strong>Операции:</strong> Поставете/махнете отметката на съответната таб-бутона, за да включите/изключите операцията при генериране. Табовете остават кликащи за настройка.</p>
        <p><strong>Деление (десетични):</strong> Новата опция "Разрешени десетични числа" позволява да се генерират примери при деление с дробни (напр. 1.2, 0.5) делители и/или частни (по избор). По подразбиране е включено.</p>
        <p><strong>Диалог за запазване:</strong> Последната използвана директория за запис се запазва и ще бъде отворена по подразбиране при следващо записване.</p>
        """
        QMessageBox.information(self, "Помощ", help_text)

    # -------------------------
    # NEW: tab change handler to change only the tabs and the area under them
    # -------------------------
    def on_tab_changed(self, index):
        color = self.custom_tab_bar.tab_colors.get(index, None)
        if color:
            # Color the small area under the tabs (above the generate button)
            self.tabs_colored_area.setStyleSheet(f"background-color: {color}; border: none;")
            # Ensure the tabs themselves keep their static colors (custom painting handles that)
            # Color the container so the area including tabs inherits the same subtle tint
            self.tabs_container.setStyleSheet(f"background-color: {color};")
            # Do not alter the overall settings group background
            self.settings_group.setStyleSheet("")
        else:
            self.tabs_colored_area.setStyleSheet("")
            self.tabs_container.setStyleSheet("")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MathWorksheetGenerator()
    window.show()
    sys.exit(app.exec_())
